# main.py

from flask import Blueprint, render_template, request
from flask_login import login_required, current_user
import glob
import docx

# set FLASK_RUN_HOST=0.0.0.0
# set FLASK_RUN_PORT=5000
# export FLASK_APP=project
# flask run --host=0.0.0.0


main = Blueprint('main', __name__)

passcode = "123"

# Sample data for blogs
posts = []

# Path to the folder containing docx files
path = "project/CFF Sample Stories/"

# Get a list of all docx files in the folder
docx_files = glob.glob(path + "*.docx")

# Loop through the files and read the heading and body of each file
print("Hello world!!")
mid = len(posts)+1
for file in docx_files:
    doc = docx.Document(file)
    heading = doc.paragraphs[0].text
    keywords = doc.paragraphs[1].text
    keywords = keywords.split(',')
    body = "\n".join([p.text for p in doc.paragraphs[2:]])
    # print("Heading: ", heading)
    # print("Body: ", body)
    print(len(doc.paragraphs))
    if mid==6:
        print("Body: ", body)
    posts.append({'id':mid, 'title': heading, 'content': body, 'keywords': keywords})
    mid = mid+1

@main.route('/')
def index():
    return render_template('index.html')

@main.route('/profile')
@login_required
def profile():
    return render_template('profile.html', name=current_user.name)

@main.route('/articles')
@login_required
def articles():
    return render_template('articles.html', name=current_user.name, posts=posts)

@main.route('/search')
@login_required
def search():
    query = request.args.get('query', '')
    search_results = [article for article in posts 
                      if (query.lower() in article['title'].lower() 
                          or query.lower() in article['content'].lower()
                          or query.lower() in " ".join(article['keywords']).lower())]
    return render_template('searchpage.html', 
                           name=current_user.name, 
                           posts=search_results, message = "Search result for '"+query+"'")


# Blog page
@main.route('/blog/<int:post_id>')
@login_required
def blog(post_id):
    post = next((p for p in posts if p['id'] == post_id), None)
    print(post["content"])
    if post:
        return render_template('blog.html', post=post)
    else:
        return render_template('error.html', error='Blog not found')